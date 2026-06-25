import sys
import os

try:
    import docx
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx dependency...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = docx.Document()
    
    # 1. Page Setup (A4 size: 8.27 x 11.69 inches)
    # Margins: Left 1.25", Right 1.0", Top 0.75", Bottom 0.75"
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.0)
    
    # Read project_report.md
    with open("project_report.md", "r", encoding="utf-8") as f:
        content = f.read()
        
    lines = content.split("\n")
    
    # Font configurator helper
    def set_font(run, font_name="Times New Roman", size_pt=12, bold=False, italic=False):
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.italic = italic
        
    in_title = True
    
    for line in lines:
        line_strip = line.strip()
        if not line_strip:
            continue
            
        # Divider marks page transitions
        if line_strip == "---":
            in_title = False
            doc.add_page_break()
            continue
            
        # Chapter and Title Headings (#)
        if line_strip.startswith("# "):
            val = line_strip[2:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.5
            
            if "Chapter" in val or val.lower() in ["abstract", "references"]:
                # Chapter Heading: 16pt bold centered
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(val)
                set_font(run, size_pt=16, bold=True)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if in_title else WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(val)
                set_font(run, size_pt=16, bold=True)
                
        # Main Headings (##)
        elif line_strip.startswith("## "):
            val = line_strip[3:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(val)
            # Main Heading: 14pt bold
            set_font(run, size_pt=14, bold=True)
            
        # Sub-headings (###)
        elif line_strip.startswith("### "):
            val = line_strip[4:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.5
            
            if in_title:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            run = p.add_run(val)
            # Sub Heading: 12pt bold
            set_font(run, size_pt=12, bold=True)
            
        # Bullet Lists
        elif line_strip.startswith("* ") or line_strip.startswith("• "):
            val = line_strip[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(val)
            set_font(run, size_pt=12)
            
        # Ordered lists or specific instructions
        elif any(line_strip.startswith(prefix) for prefix in ["1. ", "2. ", "3. ", "4. ", "5. "]):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(line_strip)
            set_font(run, size_pt=12)
            
        else:
            # Normal Paragraph: 1.5 line spacing, 12pt font
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(6)
            
            if in_title:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
            # Process inline Markdown bold formatting (**text**)
            parts = line_strip.split("**")
            for idx, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                is_bold = (idx % 2 == 1)
                set_font(run, size_pt=12, bold=is_bold)
                
    doc.save("project_report.docx")
    print("Word Document 'project_report.docx' generated successfully!")

if __name__ == "__main__":
    create_report()
